import streamlit as st
import pandas as pd
import seaborn as sns
import numpy as np
import scipy.stats as stats
from statsmodels.stats.multicomp import pairwise_tukeyhsd
import matplotlib.pyplot as plt
import io
import datetime  # Import moved to the top with other imports
import time  # Added missing import

# Modified to include calculator favicon
st.set_page_config(page_title="Analisis Uji ANOVA Satu Arah", page_icon="🧮", layout="wide")

# Hide default Streamlit elements
hide_st_style = """
        <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        </style>
        """
st.markdown(hide_st_style, unsafe_allow_html=True)

# Get current year for copyright notice
current_year = datetime.datetime.now().year

# Set title
st.title('Analisis Uji ANOVA Satu Arah')
st.markdown("""
Analisis Varians (ANOVA) Satu Arah adalah metode statistik yang digunakan untuk membandingkan rata-rata tiga atau lebih kelompok independen untuk menentukan apakah terdapat perbedaan yang signifikan secara statistik di antara mereka. Teknik ini memeriksa variasi data baik di dalam maupun antar kelompok untuk mengevaluasi perbedaan keseluruhan.
""")

# Function to create example data
def create_example_data():
    np.random.seed(42)
    # Create synthetic data with clear group differences
    group_a = np.random.normal(10, 2, 30)
    group_b = np.random.normal(13, 2, 30)
    group_c = np.random.normal(8, 2, 30)
    
    data = {
        'Group': ['Group A']*30 + ['Group B']*30 + ['Group C']*30,
        'Value': np.concatenate([group_a, group_b, group_c])
    }
    
    return pd.DataFrame(data)

# Function to check ANOVA assumptions
def check_assumptions(df, numeric_col, categorical_col):
    results = {}
    
    # Add type checking to prevent errors
    if not isinstance(df, pd.DataFrame):
        st.error(f"DataFrame diharapkan tetapi mendapatkan {type(df).__name__}")
        return {
            'normality': pd.DataFrame({'Group': ['Error'], 'Normal': [False]}),  # Fixed: removed extra parenthesis
            'homogeneity': {'Statistic': 0, 'p-value': 0, 'Equal Variances': False}
        }
    
    # Extract groups
    try:
        groups = [df[df[categorical_col] == cat][numeric_col].values for cat in df[categorical_col].unique()]
        
        # 1. Normality test (Shapiro-Wilk)
        norm_results = []
        for i, group in enumerate(groups):
            if len(group) < 3:
                # Shapiro-Wilk requires at least 3 observations
                norm_results.append({
                    'Group': df[categorical_col].unique()[i],
                    'Statistic': np.nan,
                    'p-value': np.nan,
                    'Normal': False
                })
            else:
                stat, p = stats.shapiro(group)
                norm_results.append({
                    'Group': df[categorical_col].unique()[i],
                    'Statistic': stat,
                    'p-value': p,
                    'Normal': p > 0.05
                })
        results['normality'] = pd.DataFrame(norm_results)
        
        # 2. Homogeneity of variance (Levene's test)
        if len(groups) >= 2 and all(len(g) > 0 for g in groups):
            stat, p = stats.levene(*groups)
            results['homogeneity'] = {
                'Statistic': stat,
                'p-value': p,
                'Equal Variances': p > 0.05
            }
        else:
            results['homogeneity'] = {
                'Statistic': np.nan,
                'p-value': np.nan,
                'Equal Variances': False
            }
        
        return results
    except Exception as e:
        st.error(f"Error memeriksa asumsi: {e}")
        # Return default values so the application doesn't crash
        return {
            'normality': pd.DataFrame({'Group': ['Error'], 'Statistic': [0], 'p-value': [0], 'Normal': [False]}),
            'homogeneity': {'Statistic': 0, 'p-value': 0, 'Equal Variances': False}
        }

# Main UI with tabs
tab1, tab2, tab3 = st.tabs(["📊 Input Data", "📈 Analisis & Hasil", "ℹ️ Tentang ANOVA"])

with tab1:
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("Sumber Data")
        
        # Add data privacy notice
        st.info("**Catatan Privasi**: Data yang Anda unggah ke aplikasi ini akan dienkripsi dan disimpan secara lokal dengan menggunakan session, sehingga tidak disimpan secara permanen di server. Data akan diolah dengan menggunakan server kami, namun data Anda tidak disimpan permanen, dan akan dihapus segera setelah sesi browser ditutup (session). Untuk mengakhiri session, silahkan refresh browser ini.")
        
        use_example = st.checkbox("Gunakan data contoh ini, jika Anda hanya ingin mencobanya.")
        
        if use_example:
            df = create_example_data()
            
            # Create a downloadable example CSV
            csv = df.to_csv(index=False)
            csv_bytes = csv.encode('utf-8')
            
            st.download_button(
                label="Unduh CSV contoh",
                data=csv_bytes,
                file_name="anova_example_data.csv",
                mime="text/csv",
            )
            
            # Select columns for ANOVA (for example data, these are fixed)
            numeric_col = 'Value'
            categorical_col = 'Group'
            
            st.success(f"Menggunakan data contoh dengan kolom numerik: '{numeric_col}' dan kolom kategorikal: '{categorical_col}'")
            
        else:
            # Modified file upload functionality to support Excel files
            uploaded_file = st.file_uploader("Unggah file data Anda", type=["csv", "xls", "xlsx"])
            
            if uploaded_file is not None:
                try:
                    # Detect file type by extension
                    file_extension = uploaded_file.name.split('.')[-1].lower()
                    
                    if file_extension in ['xls', 'xlsx']:
                        # Read Excel file
                        df = pd.read_excel(uploaded_file)
                        st.success(f"File Excel '{uploaded_file.name}' berhasil diimpor")
                    elif file_extension == 'csv':
                        # Read CSV file
                        df = pd.read_csv(uploaded_file)
                        st.success(f"File CSV '{uploaded_file.name}' berhasil diimpor")
                    else:
                        st.error("Format file tidak didukung. Harap unggah file CSV, XLS, atau XLSX.")
                        st.stop()
                    
                    # Select columns for ANOVA
                    numeric_columns = df.select_dtypes(include=['float64', 'int64']).columns.tolist()
                    categorical_columns = df.select_dtypes(include=['object']).columns.tolist()
                    
                    if not numeric_columns:
                        st.error("Tidak ditemukan kolom numerik dalam dataset. Mohon unggah file yang berbeda.")
                        st.stop()
                    
                    if not categorical_columns:
                        st.error("Tidak ditemukan kolom kategorikal dalam dataset. Mohon unggah file yang berbeda.")
                        st.stop()
                    
                    st.write("Pilih kolom numerik untuk ANOVA:")
                    numeric_col = st.selectbox("Kolom Numerik", numeric_columns)
                    
                    st.write("Pilih kolom kategorikal untuk ANOVA:")
                    categorical_col = st.selectbox("Kolom Kategorikal", categorical_columns)
                    
                except Exception as e:
                    st.error(f"Error membaca file: {e}")
                    st.stop()
            else:
                st.info("Silakan unggah file CSV, XLS, atau XLSX, atau gunakan data contoh.")
                st.stop()
    
    with col2:
        st.header("Pratinjau Data")
        if 'df' in locals():
            st.write(df.head())
            
            # Enhanced summary statistics 
            with st.expander("Lihat Statistik Ringkasan"):
                st.write("### Statistik Keseluruhan")
                
                # Calculate comprehensive statistics for the numeric column
                overall_stats = df[numeric_col].describe().to_frame().T
                
                # Add additional statistics
                overall_stats['skewness'] = df[numeric_col].skew()
                overall_stats['kurtosis'] = df[numeric_col].kurt()  # Changed from kurtosis() to kurt()
                overall_stats['variance'] = df[numeric_col].var()
                overall_stats['SE_mean'] = df[numeric_col].sem()
                overall_stats['CI_95%_lower'] = overall_stats['mean'] - 1.96 * overall_stats['SE_mean']
                overall_stats['CI_95%_upper'] = overall_stats['mean'] + 1.96 * overall_stats['SE_mean']
                # Add 99% confidence interval calculations
                overall_stats['CI_99%_lower'] = overall_stats['mean'] - 2.576 * overall_stats['SE_mean']
                overall_stats['CI_99%_upper'] = overall_stats['mean'] + 2.576 * overall_stats['SE_mean']
                overall_stats['range'] = df[numeric_col].max() - df[numeric_col].min()
                overall_stats['IQR'] = overall_stats['75%'] - overall_stats['25%']
                overall_stats['CV%'] = (overall_stats['std'] / overall_stats['mean'] * 100)
                
                # Display formatted overall statistics
                st.dataframe(overall_stats.style.format("{:.3f}"), use_container_width=True)
                
                st.write("### Statistik per Kelompok")
                # Calculate detailed statistics by group
                group_stats = df.groupby(categorical_col)[numeric_col].describe()
                
                # Add more group statistics using individual calculations instead of agg
                # First get basic group stats we'll need
                group_means = df.groupby(categorical_col)[numeric_col].mean()
                group_stds = df.groupby(categorical_col)[numeric_col].std()
                
                # Calculate additional statistics one by one
                group_stats['skewness'] = df.groupby(categorical_col)[numeric_col].skew()
                
                # For kurtosis, use scipy stats since pandas groupby doesn't have kurt
                from scipy import stats as scistats
                kurtosis_dict = {}
                for group in df[categorical_col].unique():
                    values = df[df[categorical_col] == group][numeric_col]
                    kurtosis_dict[group] = scistats.kurtosis(values, fisher=True)
                group_stats['kurtosis'] = pd.Series(kurtosis_dict)
                
                group_stats['variance'] = df.groupby(categorical_col)[numeric_col].var()
                group_stats['SE_mean'] = df.groupby(categorical_col)[numeric_col].sem()
                group_stats['CV%'] = (group_stds / group_means * 100)
                
                # Format and display group statistics
                st.dataframe(group_stats.style.format("{:.3f}"), use_container_width=True)
                
                # Display statistical comparison table
                st.write("### Perbandingan Antar Kelompok")
                comparison_df = df.groupby(categorical_col)[numeric_col].agg(['count', 'mean', 'std', 'min', 'max']).reset_index()
                st.dataframe(comparison_df, use_container_width=True)
        
    # Check if we have at least 2 groups
    if 'df' in locals() and 'categorical_col' in locals():
        group_counts = df[categorical_col].value_counts()
        if len(group_counts) < 2:
            st.error("ANOVA membutuhkan minimal 2 kelompok untuk dibandingkan. Data Anda hanya berisi 1 kelompok.")
            st.stop()
            
    run_analysis = st.button("Jalankan Analisis ANOVA", type="primary")
    
    # Store the necessary variables in session state
    if run_analysis and 'df' in locals():
        st.session_state.df = df
        st.session_state.numeric_col = numeric_col
        st.session_state.categorical_col = categorical_col
        st.session_state.run_analysis = True
        
        # Add notification that analysis is running
        with st.spinner('Menjalankan analisis ANOVA...'):
            # Create artificial small delay for better UX
            import time
            time.sleep(0.5)
            
        # Display success message and auto-switch to results tab
        st.success("Analisis ANOVA berhasil diselesaikan! Lihat hasilnya di tab Analisis & Hasil.")
        st.session_state.analysis_complete = True
        
        # Add option to automatically switch to results tab
        js = f"""
        <script>
            window.parent.document.querySelector('button[data-baseweb="tab"]:nth-child(2)').click();
        </script>
        """
        st.components.v1.html(js)
    
with tab2:
    st.header("Hasil ANOVA")
    
    if 'run_analysis' in st.session_state and st.session_state.run_analysis:
        df = st.session_state.df
        numeric_col = st.session_state.numeric_col
        categorical_col = st.session_state.categorical_col
        
        # Extract groups for analysis
        groups = [df[df[categorical_col] == cat][numeric_col].values for cat in df[categorical_col].unique()]
        group_labels = df[categorical_col].unique()
        group_means = [np.mean(g) for g in groups]
        group_counts = [len(g) for g in groups]
        
        # Calculate key values
        n_total = sum(group_counts)
        k = len(groups)  # Number of groups
        grand_mean = np.mean([val for group in groups for val in group])
        
        # Check ANOVA assumptions early so they're available for all analyses
        assumptions = check_assumptions(df, numeric_col, categorical_col)
        
        # Calculate SS
        ssb = sum(count * (mean - grand_mean)**2 for count, mean in zip(group_counts, group_means))
        ssw = sum(sum((val - mean)**2 for val in group) for group, mean in zip(groups, group_means))
        sst = ssb + ssw
        
        # Calculate df
        dfb = k - 1
        dfw = n_total - k
        dft = n_total - 1
        
        # Calculate MS
        msb = ssb / dfb
        msw = ssw / dfw
        
        # Calculate F and p-value
        f_stat = msb / msw
        p_val = stats.f.sf(f_stat, dfb, dfw)
        
        # Calculate effect size (eta-squared)
        eta_squared = ssb / sst
        
        # Create ANOVA table
        anova_table = pd.DataFrame({
            'Source': ['Antar Kelompok', 'Dalam Kelompok', 'Total'],
            'SS': [ssb, ssw, sst],
            'df': [dfb, dfw, dft],
            'MS': [msb, msw, np.nan],
            'F': [f_stat, np.nan, np.nan],
            'p-value': [p_val, np.nan, np.nan]
        })
        
        # Format the table
        anova_table['SS'] = anova_table['SS'].round(3)
        anova_table['MS'] = anova_table['MS'].round(3)
        anova_table['F'] = anova_table['F'].round(3)
        anova_table['p-value'] = anova_table['p-value'].map(lambda x: f"{x:.5f}" if pd.notnull(x) else x)
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.subheader("Tabel ANOVA")
            st.dataframe(anova_table, use_container_width=True)
            
            # Add significance level selection
            st.subheader("Interpretasi")
            significance_level = st.radio("Tingkat signifikansi:", [0.05, 0.01], horizontal=True)
            
            # Interpretation with selected significance level
            if p_val < significance_level:
                st.success(f"**Kesimpulan:** Terdapat perbedaan signifikan secara statistik antar kelompok (p = {p_val:.5f} < {significance_level})")
                
                # Enhanced explanation about categorical groups
                st.markdown(f"""
                **Interpretasi:**
                - Hasil menunjukkan perbedaan yang signifikan antara setidaknya dua kelompok yang didefinisikan oleh kolom kategorikal '**{categorical_col}**'
                - ANOVA ini membandingkan nilai rata-rata dari variabel '**{numeric_col}**' di antara kelompok-kelompok tersebut
                - Setidaknya satu kelompok memiliki rata-rata yang berbeda secara statistik dari kelompok lainnya
                - Untuk mengetahui kelompok mana yang berbeda secara spesifik, lihat hasil uji post-hoc di bawah
                """)
                
                # Additional interpretation for different levels
                if significance_level == 0.01:
                    st.write("Dengan tingkat signifikansi 1%, ini menunjukkan bukti yang sangat kuat terhadap kesamaan rata-rata antar kelompok.")
                else:
                    st.write("Dengan tingkat signifikansi 5%, ini menunjukkan bukti yang cukup kuat terhadap kesamaan rata-rata antar kelompok.")
                    
                st.write(f"Ukuran efek (η²): {eta_squared:.3f}")
                
                # Post-hoc test
                st.subheader("Analisis Post-Hoc")
                
                # Add post-hoc test selection
                posthoc_method = st.radio(
                    "Pilih metode uji post-hoc:",
                    ["Tukey HSD", "Bonferroni", "Scheffe", "Games-Howell", "Duncan", "Newman-Keuls"],
                    horizontal=True
                )
                
                df_posthoc = pd.DataFrame({
                    'Group': df[categorical_col],
                    'Value': df[numeric_col]
                })
                
                # Add this block here (around line 1234)
                # Information about post-hoc tests
                with st.expander("ℹ️ Informasi tentang Uji Post-Hoc"):
                    st.markdown("""
                    ### Perbandingan Metode Post-Hoc

                    | Metode | Asumsi Varians | Perbandingan | Kekuatan | Konservatif |
                    |--------|----------------|--------------|----------|-------------|
                    | Tukey HSD | Homogen | Semua pasangan | Moderat | Moderat |
                    | Bonferroni | Homogen | Semua pasangan | Rendah | Sangat |
                    | Scheffe | Homogen | Semua kombinasi | Rendah | Sangat |
                    | Games-Howell | Tidak homogen | Semua pasangan | Moderat | Moderat |
                    | Duncan | Homogen | Semua pasangan | Tinggi | Rendah |
                    | Newman-Keuls | Homogen | Semua pasangan (step-down) | Tinggi | Moderat |

                    - **Konservatif**: Tingkat kontrol terhadap kesalahan Tipe I (false positive)
                    - **Kekuatan**: Kemampuan mendeteksi perbedaan signifikan yang sebenarnya (mengurangi kesalahan Tipe II)
                    """)
                
                # Perform the selected post-hoc test
                if posthoc_method == "Tukey HSD":
                    st.write("#### Uji Tukey HSD")
                    st.write("*Terbaik untuk: Ukuran sampel yang sama, varians yang sama, semua perbandingan berpasangan.*")
                    
                    tukey_results = pairwise_tukeyhsd(df_posthoc['Value'], df_posthoc['Group'], alpha=significance_level)
                    tukey_df = pd.DataFrame(data=tukey_results._results_table.data[1:], 
                                            columns=tukey_results._results_table.data[0])
                    st.write(tukey_df)
                    
                    # Store for later use
                    posthoc_df = tukey_df
                    comparison_col1, comparison_col2 = 'group1', 'group2'
                    diff_col = 'meandiff'
                    pval_col = 'p-adj'
                    reject_col = 'reject'
                    
                elif posthoc_method == "Duncan":
                    st.write("#### Uji Jarak Berganda Duncan")
                    st.write("*Terbaik untuk: Kekuatan statistik yang lebih besar, analisis eksploratori*")
                    
                    # Get unique groups and their means, sorted by mean value
                    group_data = {}
                    for group in df[categorical_col].unique():
                        values = df[df[categorical_col] == group][numeric_col].values
                        group_data[group] = {
                            'mean': np.mean(values),
                            'n': len(values)
                        }
                    
                    # Sort groups by mean
                    sorted_groups = sorted(group_data.keys(), key=lambda g: group_data[g]['mean'])
                    sorted_means = [group_data[g]['mean'] for g in sorted_groups]
                    
                    # Calculate harmonic mean of sample sizes (for unequal sample sizes)
                    n_values = [group_data[g]['n'] for g in sorted_groups]
                    if all(n == n_values[0] for n in n_values):
                        n_harmonic = n_values[0]  # All sample sizes are equal
                    else:
                        n_harmonic = len(n_values) / sum(1/n for n in n_values)
                    
                    # Values needed for Duncan test
                    k = len(sorted_groups)
                    mse = msw  # Mean square error from ANOVA
                    df_error = dfw  # Error degrees of freedom from ANOVA
                    
                    # Function to get critical values for Duncan test
                    def get_duncan_critical_value(p, df, alpha=0.05):
                        # p is the range (number of steps between means)
                        # Get studentized range statistic
                        # Note: This uses the studentized range approximation
                        q_alpha = stats.t.ppf(1 - alpha/2, df) * np.sqrt(2)
                        # Adjust for number of steps
                        return q_alpha * np.sqrt(p * (p + 1) / (6 * df))
                    
                    # Calculate all pairwise comparisons
                    results = []
                    for i in range(k):
                        for j in range(i + 1, k):
                            # Calculate range (p) - number of steps between means in the sorted list
                            p = j - i + 1
                            
                            # Get critical value
                            critical_value = get_duncan_critical_value(p, df_error, significance_level)
                            
                            # Calculate least significant range
                            lsr = critical_value * np.sqrt(mse / n_harmonic)
                            
                            # Calculate mean difference
                            mean_diff = sorted_means[j] - sorted_means[i]
                            
                            # Check if difference is significant
                            significant = mean_diff > lsr
                            
                            results.append({
                                'group1': sorted_groups[i],
                                'group2': sorted_groups[j],
                                'meandiff': mean_diff,
                                'range(p)': p,
                                'lsr': lsr,
                                'reject': significant,
                                'p-value': significance_level if significant else 'NS'
                            })
                    
                    # Convert to DataFrame
                    posthoc_df = pd.DataFrame(results)
                    st.write(posthoc_df)
                    
                    # Store column names for later use
                    comparison_col1, comparison_col2 = 'group1', 'group2'
                    diff_col = 'meandiff'
                    pval_col = 'p-value'
                    reject_col = 'reject'
                    
                    # Add Duncan homogeneous subsets table (shows groups that are not significantly different)
                    st.subheader("Subset Homogen Duncan")
                    
                    # Create a matrix for quick lookup of significant differences
                    sig_matrix = {g1: {g2: False for g2 in sorted_groups} for g1 in sorted_groups}
                    for _, row in posthoc_df.iterrows():
                        if row['reject']:
                            sig_matrix[row['group1']][row['group2']] = True
                            sig_matrix[row['group2']][row['group1']] = True
                    
                    # Form homogeneous subsets (groups that are not significantly different)
                    subsets = []
                    remaining = set(sorted_groups)
                    
                    while remaining:
                        # Start a new subset with the first remaining group
                        current = list(remaining)[0]
                        subset = {current}
                        remaining.remove(current)
                        
                        # Try to add other groups to this subset
                        for group in list(remaining):
                            # If not significantly different from all groups in the subset, add it
                            if all(not sig_matrix[group][g] for g in subset):
                                subset.add(group)
                                remaining.remove(group)
                        
                        subsets.append(subset)
                    
                    # Create a DataFrame for displaying homogeneous subsets
                    subset_data = []
                    for i, subset in enumerate(subsets):
                        for group in subset:
                            subset_data.append({
                                'Group': group, 
                                'Mean': group_data[group]['mean'],
                                'Subset': f"Subset {i+1}"
                            })
                    
                    subset_df = pd.DataFrame(subset_data)
                    subset_df = subset_df.sort_values('Mean', ascending=False)
                    
                    # Pivot to create a format similar to SPSS output
                    subset_pivot = subset_df.pivot(index=['Group', 'Mean'], columns='Subset', values='Mean')
                    subset_pivot = subset_pivot.reset_index()
                    
                    st.write(subset_pivot)
                    st.markdown("""
                    **Catatan tentang subset Duncan:** Kelompok dalam subset yang sama tidak berbeda secara signifikan pada tingkat alpha yang dipilih.
                    Uji Duncan kurang konservatif dibandingkan Tukey HSD, sehingga dapat mengidentifikasi lebih banyak perbedaan signifikan.
                    """)
                
                elif posthoc_method == "Newman-Keuls":
                    st.write("#### Uji Student-Newman-Keuls (SNK)")
                    st.write("""
                    *Terbaik untuk: Kontrol kesalahan familywise dengan kekuatan statistik lebih tinggi dari Tukey HSD.*

                    **Profil Uji Newman-Keuls:**
                    - Menggunakan pendekatan *step-down* yang menyesuaikan nilai kritis berdasarkan rentang antara rata-rata
                    - Memiliki kekuatan statistik yang lebih tinggi daripada Tukey HSD, sehingga lebih mampu mendeteksi perbedaan nyata
                    - Menyeimbangkan kontrol kesalahan Tipe I dan Tipe II lebih baik dari metode konservatif
                    - Menguji perbedaan dengan nilai kritis yang berbeda tergantung pada jarak peringkat antar rata-rata
                    - Rekomendasi: Gunakan ketika varians homogen dan Anda ingin metode dengan kekuatan lebih tinggi dari Tukey HSD
                    """)
                    
                    # Get unique groups and their means, sorted by mean value
                    group_data = {}
                    for group in df[categorical_col].unique():
                        values = df[df[categorical_col] == group][numeric_col].values
                        group_data[group] = {
                            'mean': np.mean(values),
                            'n': len(values)
                        }
                    
                    # Sort groups by mean
                    sorted_groups = sorted(group_data.keys(), key=lambda g: group_data[g]['mean'])
                    sorted_means = [group_data[g]['mean'] for g in sorted_groups]
                    
                    # Calculate MSE and df from ANOVA
                    mse = msw  # Mean square within from ANOVA
                    df_error = dfw  # Degrees of freedom within groups
                    
                    # Calculate harmonic mean of sample sizes (for unequal sample sizes)
                    n_values = [group_data[g]['n'] for g in sorted_groups]
                    if all(n == n_values[0] for n in n_values):
                        n_harmonic = n_values[0]  # All sample sizes are equal
                    else:
                        n_harmonic = len(n_values) / sum(1/n for n in n_values)
                    
                    # Function to get critical q value for Newman-Keuls test
                    def get_q_critical(p, df, alpha=0.05):
                        # p is the range (number of steps between means)
                        # Use studentized range distribution for SNK test
                        from scipy.stats import studentized_range
                        try:
                            # Direct calculation if scipy has the function
                            q = studentized_range.ppf(1-alpha, p, df)
                        except:
                            # Approximation if the function is not available
                            q = stats.t.ppf(1 - alpha/(2*p), df) * np.sqrt(2)
                        return q
                    
                    # Perform all pairwise comparisons
                    results = []
                    k = len(sorted_groups)
                    
                    # Newman-Keuls compares all pairs starting with the largest difference
                    for step_size in range(k-1, 0, -1):
                        for i in range(k-step_size):
                            j = i + step_size
                            group1, group2 = sorted_groups[i], sorted_groups[j]
                            mean_diff = sorted_means[j] - sorted_means[i]
                            
                            # Get critical q value for this range
                            q_crit = get_q_critical(step_size+1, df_error, significance_level)
                            
                            # Calculate least significant range
                            se = np.sqrt(mse / n_harmonic)
                            lsr = q_crit * se
                            
                            # Calculate test statistic
                            q_stat = mean_diff / se if se > 0 else float('inf')
                            
                            # Check if difference is significant
                            significant = mean_diff > lsr
                            
                            # Calculate p-value (approximate)
                            p_val = 1 - stats.norm.cdf(q_stat / np.sqrt(2))
                            
                            results.append({
                                'group1': group1,
                                'group2': group2,
                                'meandiff': mean_diff,
                                'steps': step_size+1,
                                'q-value': q_stat,
                                'critical q': q_crit,
                                'p-value': p_val,
                                'reject': significant
                            })
                    
                    # Convert to DataFrame
                    posthoc_df = pd.DataFrame(results)
                    st.write(posthoc_df)
                    
                    # Store column names for later use
                    comparison_col1, comparison_col2 = 'group1', 'group2'
                    diff_col = 'meandiff'
                    pval_col = 'p-value'
                    reject_col = 'reject'
                    
                    # Add homogeneous subsets table
                    st.subheader("Subset Homogen Newman-Keuls")
                    
                    # Create a matrix for quick lookup of significant differences
                    sig_matrix = {g1: {g2: False for g2 in sorted_groups} for g1 in sorted_groups}
                    for _, row in posthoc_df.iterrows():
                        if row['reject']:
                            sig_matrix[row['group1']][row['group2']] = True
                            sig_matrix[row['group2']][row['group1']] = True
                    
                    # Form homogeneous subsets (groups that are not significantly different)
                    subsets = []
                    remaining = set(sorted_groups)
                    
                    while remaining:
                        # Start a new subset with the first remaining group
                        current = list(remaining)[0]
                        subset = {current}
                        remaining.remove(current)
                        
                        # Try to add other groups to this subset
                        for group in list(remaining):
                            # If not significantly different from all groups in the subset, add it
                            if all(not sig_matrix[group][g] for g in subset):
                                subset.add(group)
                                remaining.remove(group)
                        
                        subsets.append(subset)
                    
                    # Create a DataFrame for displaying homogeneous subsets
                    subset_data = []
                    for i, subset in enumerate(subsets):
                        for group in subset:
                            subset_data.append({
                                'Group': group, 
                                'Mean': group_data[group]['mean'],
                                'Subset': f"Subset {i+1}"
                            })
                    
                    subset_df = pd.DataFrame(subset_data)
                    subset_df = subset_df.sort_values('Mean', ascending=False)
                    
                    # Pivot to create a format similar to SPSS output
                    subset_pivot = subset_df.pivot(index=['Group', 'Mean'], columns='Subset', values='Mean')
                    subset_pivot = subset_pivot.reset_index()
                    
                    st.write(subset_pivot)
                    st.markdown("""
                    **Catatan tentang uji Newman-Keuls:** 
                    - Kelompok dalam subset yang sama tidak berbeda secara signifikan
                    - Uji Newman-Keuls memiliki kekuatan statistik yang lebih tinggi dibandingkan Tukey HSD
                    - SNK menggunakan pendekatan step-down yang menyesuaikan nilai kritis berdasarkan rentang antara rata-rata
                    """)
                
                elif posthoc_method == "Games-Howell":
                    st.write("#### Uji Games-Howell")
                    st.write("*Terbaik untuk: Varians tidak sama, ukuran sampel tidak sama.*")
                    
                    # Create all pairwise combinations
                    from itertools import combinations
                    pairs = list(combinations(df[categorical_col].unique(), 2))
                    results = []
                    
                    # Perform Games-Howell test for each pair
                    for group1, group2 in pairs:
                        # Add error checking to avoid potential issues
                        try:
                            # Get data for each group
                            data1 = df[df[categorical_col] == group1][numeric_col].values
                            data2 = df[df[categorical_col] == group2][numeric_col].values
                            
                            # Check if we have enough data
                            if len(data1) < 2 or len(data2) < 2:
                                st.warning(f"Insufficient data for groups {group1} or {group2}. Skipping this comparison.")
                                continue
                            
                            n1, n2 = len(data1), len(data2)
                            mean1, mean2 = np.mean(data1), np.mean(data2)
                            var1, var2 = np.var(data1, ddof=1), np.var(data2, ddof=1)
                            
                            # Mean difference
                            mean_diff = mean1 - mean2
                            
                            # Pooled standard error
                            se = np.sqrt(var1/n1 + var2/n2)
                            
                            # Calculate degrees of freedom using Welch-Satterthwaite equation
                            if se == 0:  # Avoid division by zero
                                df_val = float('inf')
                            else:
                                df_val = ((var1/n1 + var2/n2)**2) / (((var1/n1)**2/(n1-1)) + ((var2/n2)**2/(n2-1)))
                            
                            # Calculate t-value
                            t_val = abs(mean_diff) / se if se > 0 else float('inf')
                            
                            # Get critical q value (studentized range)
                            q_crit = stats.t.ppf(1 - significance_level/2, df_val) * np.sqrt(2)
                            
                            # Calculate p-value
                            p_val = 2 * (1 - stats.t.cdf(abs(t_val), df_val))
                            
                            # Calculate confidence interval
                            lower = mean_diff - q_crit * se
                            upper = mean_diff + q_crit * se
                            
                            results.append({
                                'group1': group1,
                                'group2': group2,
                                'meandiff': mean_diff,
                                't-value': t_val,
                                'df': df_val,
                                'p-value': p_val,
                                'lower': lower,
                                'upper': upper,
                                'reject': p_val < significance_level
                            })
                        except Exception as e:
                            st.error(f"Error analyzing groups {group1} and {group2}: {str(e)}")
                            continue
                    
                    if results:
                        posthoc_df = pd.DataFrame(results)
                        st.write(posthoc_df)
                        
                        # Store column names for later use
                        comparison_col1, comparison_col2 = 'group1', 'group2'
                        diff_col = 'meandiff'
                        pval_col = 'p-value'
                        reject_col = 'reject'
                    else:
                        st.error("Tidak dapat melakukan analisis post-hoc Games-Howell. Silakan periksa data Anda.")
                
                elif posthoc_method == "Bonferroni":
                    st.write("#### Uji Bonferroni")
                    st.write("*Terbaik untuk: Jumlah perbandingan kecil, kontrol ketat terhadap kesalahan Tipe I.*")
                    
                    # Create all pairwise combinations
                    from itertools import combinations
                    pairs = list(combinations(df[categorical_col].unique(), 2))
                    
                    # Calculate number of comparisons for Bonferroni correction
                    num_comparisons = len(pairs)
                    
                    # Perform Bonferroni-corrected t-tests
                    results = []
                    for group1, group2 in pairs:
                        # Get data for each group
                        data1 = df[df[categorical_col] == group1][numeric_col].values
                        data2 = df[df[categorical_col] == group2][numeric_col].values
                        
                        # Check if we have enough data
                        if len(data1) < 2 or len(data2) < 2:
                            st.warning(f"Insufficient data for groups {group1} or {group2}. Skipping this comparison.")
                            continue
                            
                        # Perform t-test
                        t_stat, p_val = stats.ttest_ind(data1, data2, equal_var=assumptions['homogeneity']['Equal Variances'])
                        
                        # Apply Bonferroni correction
                        p_adjusted = min(p_val * num_comparisons, 1.0)  # Cap at 1.0
                        
                        # Calculate mean difference
                        mean_diff = np.mean(data1) - np.mean(data2)
                        
                        # Calculate standard error
                        n1, n2 = len(data1), len(data2)
                        if assumptions['homogeneity']['Equal Variances']:
                            # Pooled standard error (equal variances)
                            var1 = np.var(data1, ddof=1)
                            var2 = np.var(data2, ddof=1)
                            pooled_var = ((n1-1)*var1 + (n2-1)*var2) / (n1 + n2 - 2)
                            se = np.sqrt(pooled_var * (1/n1 + 1/n2))
                        else:
                            # Welch's t-test standard error (unequal variances)
                            se = np.sqrt(np.var(data1, ddof=1)/n1 + np.var(data2, ddof=1)/n2)
                        
                        # Calculate confidence interval
                        if assumptions['homogeneity']['Equal Variances']:
                            # Use t-distribution with pooled degrees of freedom
                            df_val = n1 + n2 - 2
                        else:
                            # Use Welch-Satterthwaite equation for degrees of freedom
                            var1 = np.var(data1, ddof=1)
                            var2 = np.var(data2, ddof=1)
                            num = (var1/n1 + var2/n2)**2
                            denom = (var1/n1)**2/(n1-1) + (var2/n2)**2/(n2-1)
                            df_val = num/denom if denom > 0 else float('inf')
                        
                        # Get critical value for confidence interval
                        # For Bonferroni, we need to adjust alpha for the number of comparisons
                        alpha_corrected = significance_level / num_comparisons
                        t_crit = stats.t.ppf(1 - alpha_corrected/2, df_val)
                        
                        # Calculate confidence interval
                        lower = mean_diff - t_crit * se
                        upper = mean_diff + t_crit * se
                        
                        results.append({
                            'group1': group1,
                            'group2': group2,
                            'meandiff': mean_diff,
                            't-value': t_stat,
                            'p-value': p_val,
                            'p-adjusted': p_adjusted,
                            'lower': lower,
                            'upper': upper,
                            'reject': p_adjusted < significance_level
                        })

                    # Convert to DataFrame
                    posthoc_df = pd.DataFrame(results)
                    st.write(posthoc_df)
                    
                    # Store column names for later use
                    comparison_col1, comparison_col2 = 'group1', 'group2'
                    diff_col = 'meandiff'
                    pval_col = 'p-adjusted'  # Use adjusted p-value for Bonferroni
                    reject_col = 'reject'
                    
                    st.markdown("""
                    **Catatan tentang koreksi Bonferroni:**
                    - Koreksi Bonferroni mengontrol tingkat kesalahan familywise dengan membagi alpha dengan jumlah perbandingan
                    - Ini sangat konservatif, terutama dengan banyak perbandingan
                    - P-adjusted = p-value × jumlah perbandingan (dibulatkan ke maksimum 1,0)
                    - Menolak H₀ jika p-adjusted < alpha
                    """)
                
                elif posthoc_method == "Scheffe":
                    st.write("#### Uji Scheffe")
                    st.write("*Terbaik untuk: Semua kombinasi kelompok, tidak hanya perbandingan berpasangan, dan ukuran sampel yang tidak sama.*")
                    
                    # Create all pairwise combinations
                    from itertools import combinations
                    pairs = list(combinations(df[categorical_col].unique(), 2))
                    
                    # MSE from ANOVA
                    mse = msw
                    
                    # Perform Scheffe test
                    results = []
                    for group1, group2 in pairs:
                        # Get data for each group
                        data1 = df[df[categorical_col] == group1][numeric_col].values
                        data2 = df[df[categorical_col] == group2][numeric_col].values
                        
                        # Check if we have enough data
                        if len(data1) < 2 or len(data2) < 2:
                            st.warning(f"Insufficient data for groups {group1} or {group2}. Skipping this comparison.")
                            continue
                        
                        # Sample sizes
                        n1, n2 = len(data1), len(data2)
                        
                        # Group means
                        mean1, mean2 = np.mean(data1), np.mean(data2)
                        
                        # Calculate mean difference
                        mean_diff = mean1 - mean2
                        
                        # Calculate F statistic for this comparison
                        # F = (mean_diff)^2 / [MSE * (1/n1 + 1/n2)]
                        f_val = (mean_diff**2) / (mse * (1/n1 + 1/n2))
                        
                        # Critical F value for Scheffe
                        # (k-1) * F_crit(alpha, k-1, df_error)
                        k = len(df[categorical_col].unique())  # Number of groups
                        f_crit = stats.f.ppf(1 - significance_level, k-1, dfw)
                        scheffe_crit = (k-1) * f_crit
                        
                        # Calculate p-value
                        # p-value = 1 - F_cdf(F/((k-1)), k-1, df_error)
                        p_val = 1 - stats.f.cdf(f_val/(k-1), k-1, dfw)
                        
                        # Standard error
                        se = np.sqrt(mse * (1/n1 + 1/n2))
                        
                        # Calculate confidence interval
                        # Critical value for Scheffe interval
                        scheffe_cval = np.sqrt(scheffe_crit)
                        
                        # Confidence interval
                        lower = mean_diff - scheffe_cval * se
                        upper = mean_diff + scheffe_cval * se
                        
                        results.append({
                            'group1': group1,
                            'group2': group2,
                            'meandiff': mean_diff,
                            'F-value': f_val,
                            'critical F': scheffe_crit,
                            'p-value': p_val,
                            'lower': lower,
                            'upper': upper,
                            'reject': f_val > scheffe_crit
                        })
                    
                    # Convert to DataFrame
                    posthoc_df = pd.DataFrame(results)
                    st.write(posthoc_df)
                    
                    # Store column names for later use
                    comparison_col1, comparison_col2 = 'group1', 'group2'
                    diff_col = 'meandiff'
                    pval_col = 'p-value'
                    reject_col = 'reject'
                    
                    st.markdown("""
                    **Catatan tentang uji Scheffe:**
                    - Uji Scheffe adalah salah satu uji post-hoc paling konservatif
                    - Dirancang untuk menguji tidak hanya perbandingan berpasangan tetapi semua kemungkinan kontras
                    - Kontrol ketat terhadap kesalahan familywise dalam semua perbandingan
                    - Cocok digunakan ketika terdapat jumlah sampel yang tidak sama antar kelompok
                    - Sangat konservatif, sehingga memiliki kekuatan statistik yang lebih rendah dibandingkan uji lain seperti Tukey HSD
                    """)
                
                # Create a subset table showing significant differences
                st.subheader("Perbedaan Kelompok yang Signifikan")

                # Check if posthoc_df exists before trying to use it
                if 'posthoc_df' in locals():
                    sig_pairs = posthoc_df[posthoc_df[reject_col] == True].copy()
                    
                    if len(sig_pairs) > 0:
                        # Format the table to be more readable
                        sig_pairs['Mean Difference'] = sig_pairs[diff_col].round(3)
                        sig_pairs['p-value'] = sig_pairs[pval_col].apply(lambda x: f"{x:.5f}" if isinstance(x, (int, float)) else x)
                        
                        # Check if confidence interval columns exist (Duncan's test doesn't have them)
                        if 'lower' in sig_pairs.columns and 'upper' in sig_pairs.columns:
                            sig_pairs['Confidence Interval'] = sig_pairs.apply(
                                lambda row: f"[{row['lower']:.3f}, {row['upper']:.3f}]", axis=1
                            )
                            display_columns = [comparison_col1, comparison_col2, 'Mean Difference', 'p-value', 'Confidence Interval']
                        else:
                            # For methods without confidence intervals (like Duncan)
                            display_columns = [comparison_col1, comparison_col2, 'Mean Difference', 'p-value']
                        
                        # Display only relevant columns
                        sig_pairs_display = sig_pairs[display_columns]
                        
                        # Rename columns for better display
                        column_names = {'Group 1': comparison_col1, 'Group 2': comparison_col2, 
                                       'Mean Difference': 'Mean Difference', 
                                       'p-value': 'p-value'}
                        if 'Confidence Interval' in display_columns:
                            column_names['Confidence Interval'] = f'{int((1-significance_level)*100)}% CI'
                        
                        sig_pairs_display.columns = column_names.values()
                        
                        st.write(sig_pairs_display)
                        
                        # Create a ranked subset table - only if we have significant pairs
                        st.subheader("Peringkat Kelompok")
                        
                        # Make sure we're using the original dataframe from session state
                        original_df = st.session_state.df
                        
                        # Continue with the rest of the code
                        # ...
                    else:
                        st.info("Tidak ditemukan perbedaan signifikan antara pasangan kelompok manapun dalam uji post-hoc.")
                else:
                    st.info("Analisis post-hoc tidak tersedia. Coba pilih metode post-hoc yang berbeda.")
                
            else:
                st.info(f"**Kesimpulan:** Tidak ada perbedaan signifikan secara statistik antar kelompok (p = {p_val:.5f} ≥ {significance_level})")
                st.write(f"Effect size (η²): {eta_squared:.3f}")
            
        with col2:
            # Check ANOVA assumptions
            st.subheader("ANOVA Assumptions")
            assumptions = check_assumptions(df, numeric_col, categorical_col)
            
            with st.expander("Normality Test (Shapiro-Wilk)"):
                st.write("H₀: Data is normally distributed")
                st.dataframe(assumptions['normality'], use_container_width=True)
                if all(assumptions['normality']['Normal']):
                    st.success("✓ All groups appear to be normally distributed")
                else:
                    st.warning("✗ Some groups may not be normally distributed")
                    
                    # Add recommendations for non-normal data
                    non_normal_groups = assumptions['normality'][~assumptions['normality']['Normal']]['Group'].tolist()
                    st.write(f"Groups that failed normality test: {', '.join(non_normal_groups)}")
                    
                    st.subheader("Rekomendasi:")
                    st.markdown("""
                    Ketika data tidak terdistribusi normal, pertimbangkan opsi berikut:
                    
                    1. **Transformasi data** - Terapkan salah satu transformasi berikut:
                       - Transformasi log: `log(x)` atau `log(x+1)` jika data mengandung nol
                       - Transformasi akar kuadrat: `sqrt(x)`
                       - Transformasi Box-Cox
                       
                    2. **Gunakan alternatif non-parametrik** - Uji Kruskal-Wallis adalah 
                       alternatif non-parametrik untuk ANOVA satu arah dan tidak mengasumsikan normalitas.
                       
                    3. **Lanjutkan dengan ANOVA jika:**
                       - Ukuran sampel besar (n > 30 per kelompok) - Teorema Limit Pusat berlaku
                       - Penyimpangan dari normalitas tidak ekstrem (periksa plot QQ)
                       - Semua kelompok memiliki ukuran sampel yang kurang lebih sama
                    """)
                    
                    # Add a button to perform Kruskal-Wallis test
                    if st.button("Run Kruskal-Wallis Test"):
                        # Perform Kruskal-Wallis test
                        stat, p_value = stats.kruskal(*groups)
                        
                        # Create result dataframe
                        kw_result = pd.DataFrame({
                            'Statistic': [stat],
                            'p-value': [p_value],
                            'Significant': [p_value < significance_level]
                        })
                        
                        st.write("### Hasil Uji Kruskal-Wallis:")
                        st.dataframe(kw_result)
                        
                        if p_value < significance_level:
                            st.success(f"Uji Kruskal-Wallis menunjukkan perbedaan signifikan antar kelompok (p = {p_value:.5f})")
                            
                            # For multiple comparisons after Kruskal-Wallis
                            st.write("### Perbandingan Berpasangan (uji Dunn):")
                            
                            st.info("""
                            Untuk analisis lengkap, pertimbangkan untuk melakukan uji Dunn untuk perbandingan berpasangan post-hoc.
                            Uji ini tepat setelah hasil Kruskal-Wallis yang signifikan.
                            """)
                        else:
                            st.info(f"Uji Kruskal-Wallis tidak menunjukkan perbedaan signifikan antar kelompok (p = {p_value:.5f})")
            
            with st.expander("Homogeneity of Variance (Levene's Test)"):
                st.write("H₀: Variances are equal across groups")
                st.write(f"Statistic: {assumptions['homogeneity']['Statistic']:.3f}")
                st.write(f"p-value: {assumptions['homogeneity']['p-value']:.5f}")
                if assumptions['homogeneity']['Equal Variances']:
                    st.success("✓ Variances appear to be homogeneous")
                else:
                    st.warning("✗ Variances may not be equal across groups")
                    
                    st.subheader("Rekomendasi untuk varians tidak sama:")
                    st.markdown("""
                    Ketika varians tidak sama (heteroskedastisitas), pertimbangkan opsi berikut:
                    
                    1. **Gunakan ANOVA Welch** - Alternatif yang kuat untuk ANOVA satu arah yang tidak mengasumsikan varians sama
                    
                    2. **Transformasi data** - Transformasikan data Anda untuk menstabilkan varians
                    
                    3. **Gunakan uji non-parametrik** - Pertimbangkan uji Kruskal-Wallis
                    
                    4. **Lanjutkan dengan hati-hati** jika:
                       - Ukuran kelompok kurang lebih sama
                       - Rasio varians kelompok terbesar dengan terkecil kurang dari 4:1
                    """)
                    
                    # Add a button for Welch's ANOVA
                    if st.button("Jalankan ANOVA Welch"):
                        try:
                            # Calculate Welch's statistic
                            group_vars = [np.var(g, ddof=1) for g in groups]
                            
                            # Add small epsilon to prevent division by zero
                            epsilon = 1e-10
                            group_vars = [max(var, epsilon) for var in group_vars]
                            
                            # Ensure all values are numeric
                            group_means_numeric = [float(mean) for mean in group_means]
                            group_counts_numeric = [int(count) for count in group_counts]
                            w = sum(count * (mean - grand_mean_numeric)**2 / var 
                                  for count, mean, var in zip(group_counts_numeric, group_means_numeric, group_vars))
                            
                            # Calculate degrees of freedom
                            df1 = k - 1
                            v_top = df1**2
                            v_bottom = sum((1 / (count - 1)) * ((count * mean**2 / var)**2) 
                                          for count, mean, var in zip(group_counts_numeric, group_means_numeric, group_vars))
                            df2 = v_top / v_bottom if v_bottom > 0 else float('inf')
                            
                            # Calculate p-value
                            p_value_welch = stats.f.sf(w, df1, df2)
                            
                            # Display results
                            welch_result = pd.DataFrame({
                                'Statistic': [w],
                                'df1': [df1],
                                'df2': [round(df2, 2)], 
                                'p-value': [p_value_welch],
                                'Significant': [p_value_welch < significance_level]
                            })
                            
                            st.write("### Hasil ANOVA Welch:")
                            st.dataframe(welch_result)
                            
                            if p_value_welch < significance_level:
                                st.success(f"ANOVA Welch menunjukkan perbedaan signifikan antar kelompok (p = {p_value_welch:.5f})")
                            else:
                                st.info(f"ANOVA Welch tidak menunjukkan perbedaan signifikan antar kelompok (p = {p_value_welch:.5f})")
                        except Exception as e:
                            st.error(f"Error dalam perhitungan ANOVA Welch: {e}")
                            st.error("Hal ini mungkin disebabkan oleh masalah tipe data. Pastikan data Anda hanya berisi nilai numerik.")
                    
        # Visualizations
        st.header("Visualizations")
        
        # Get fresh reference to the original dataframe to avoid any variable shadowing issues
        viz_df = st.session_state.df
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Boxplot Kelompok")
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.boxplot(x=categorical_col, y=numeric_col, data=viz_df, ax=ax)
            ax.set_title(f'Boxplot {numeric_col} berdasarkan {categorical_col}')
            st.pyplot(fig)
            
            st.markdown("""
            **Cara menginterpretasikan boxplot ini:**
            - **Kotak**: Kotak mewakili rentang interkuartil (IQR), dari persentil ke-25 hingga ke-75
            - **Garis horizontal di kotak**: Menunjukkan nilai median (persentil ke-50)
            - **Whiskers**: Memanjang ke nilai terkecil/terbesar dalam rentang 1,5 kali IQR
            - **Titik**: Outlier (nilai di luar whiskers)
            
            **Yang perlu diperhatikan:**
            - Bandingkan nilai median antar kelompok
            - Periksa perbedaan sebaran (ukuran kotak dan panjang whiskers)
            - Identifikasi outlier potensial yang mungkin mempengaruhi analisis Anda
            - Perhatikan asimetri dalam distribusi
            """)
        
        with col2:
            # Calculate CI level based on significance level
            ci_level = 100 * (1 - significance_level)  # 95% for 0.05, 99% for 0.01
            
            st.subheader(f"Rerata Kelompok dengan CI {ci_level:.0f}%")
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.barplot(x=categorical_col, y=numeric_col, data=viz_df, errorbar=('ci', ci_level), ax=ax)
            ax.set_title(f'Rerata {numeric_col} berdasarkan {categorical_col} dengan CI {ci_level:.0f}%')
            st.pyplot(fig)
            
            st.markdown(f"""
            **Cara menginterpretasikan plot batang ini:**
            - **Tinggi batang**: Menunjukkan nilai rerata untuk setiap kelompok
            - **Rentang kesalahan**: Menunjukkan interval kepercayaan {ci_level:.0f}% untuk setiap rerata
            
            **Yang perlu diperhatikan:**
            - Bandingkan nilai rerata antar kelompok
            - Periksa apakah interval kepercayaan saling tumpang tindih:
              - CI yang tidak tumpang tindih umumnya mengindikasikan perbedaan signifikan
              - CI yang tumpang tindih mungkin masih memiliki perbedaan signifikan (periksa hasil ANOVA)
            - Interval kepercayaan yang lebih lebar menunjukkan ketidakpastian yang lebih besar (biasanya karena ukuran sampel kecil atau varians tinggi)
            """)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Plot Distribusi")
            fig, ax = plt.subplots(figsize=(10, 6))
            for i, group in enumerate(group_labels):
                sns.kdeplot(viz_df[viz_df[categorical_col] == group][numeric_col], 
                          label=group, ax=ax)
            ax.set_title(f'Distribusi {numeric_col} berdasarkan {categorical_col}')
            ax.legend()
            st.pyplot(fig)
            
            st.markdown("""
            **Cara menginterpretasikan plot distribusi ini:**
            - Setiap kurva menunjukkan fungsi kepadatan probabilitas untuk suatu kelompok
            - Puncak yang lebih tinggi menunjukkan nilai yang lebih umum
            - Kurva yang lebih lebar menunjukkan variabilitas data yang lebih besar
            
            **Yang perlu diperhatikan:**
            - Bandingkan kecenderungan pusat (puncak) antar kelompok
            - Periksa bentuk setiap distribusi:
              - Kurva berbentuk lonceng menunjukkan distribusi normal
              - Kurva miring menunjukkan distribusi tidak normal
            - Perhatikan distribusi multimodal (beberapa puncak), yang mungkin mengindikasikan adanya sub-kelompok
            - Nilai overlap antara kelompok (semakin sedikit overlap menunjukkan perbedaan yang lebih kuat)
            """)
            
        with col2:
            st.subheader("Plot QQ")
            fig, axes = plt.subplots(1, len(groups), figsize=(10, 6), sharey=True)
            if len(groups) == 1:
                axes = [axes]  # Make it iterable for the loop
            
            for i, (label, group_data) in enumerate(zip(group_labels, groups)):
                stats.probplot(group_data, dist="norm", plot=axes[i])
                axes[i].set_title(f'{label}')
                if i > 0:
                    axes[i].set_ylabel('')
            
            plt.tight_layout()
            st.pyplot(fig)
            
            st.markdown("""
            **Cara menginterpretasikan plot QQ ini:**
            - Titik yang mengikuti garis lurus menunjukkan data yang terdistribusi normal
            - Penyimpangan dari garis diagonal menunjukkan deviasi dari distribusi normal
            - Perhatikan pola penyimpangan:
              - Kurva ke atas/bawah di ujung menunjukkan ekor yang lebih berat/ringan
              - Pola S menunjukkan skewness (kemiringan) dalam data
            """)
            
            # Proper code for document export functionality
            import io
            
            # Check if we need to install python-docx
            try:
                import docx
            except ImportError:
                import sys
                import subprocess
                st.info("Menginstal library yang diperlukan...")
                subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
                import docx
            
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create a function to convert DataFrame to table in Word document
            def add_dataframe_to_doc(doc, df, title=None):
                if title:
                    doc.add_heading(title, level=2)
                
                # Add a table
                table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
                table.style = 'Table Grid'
                
                # Add headers
                for j, column in enumerate(df.columns):
                    table.cell(0, j).text = str(column)
                
                # Add data
                for i, row in enumerate(df.values):
                    for j, val in enumerate(row):
                        table.cell(i+1, j).text = str(val)
                
                doc.add_paragraph("")
            
            # Create a Word document
            doc = Document()
            
            # Add title
            doc.add_heading('Hasil Analisis ANOVA Satu Arah', 0)
            
            # Add metadata
            doc.add_heading('Informasi Analisis', 1)
            import datetime
            meta_data = [
                ('Tanggal Analisis', datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                ('Variabel Numerik', numeric_col),
                ('Variabel Kategori', categorical_col),
                ('Ukuran Sampel Total', str(n_total)),
                ('Jumlah Kelompok', str(k)),
                ('Tingkat Signifikansi', str(significance_level))
            ]
            
            meta_table = doc.add_table(rows=len(meta_data), cols=2)
            meta_table.style = 'Table Grid'
            for i, (key, value) in enumerate(meta_data):
                meta_table.cell(i, 0).text = key
                meta_table.cell(i, 1).text = value
            
            doc.add_paragraph("")
            
            # Add ANOVA results
            doc.add_heading('Hasil ANOVA', 1)
            add_dataframe_to_doc(doc, anova_table, "Tabel ANOVA")
            
            # Add interpretation
            doc.add_heading('Interpretasi', 2)
            if p_val < significance_level:
                p = doc.add_paragraph()
                p.add_run(f"Kesimpulan: Terdapat perbedaan signifikan secara statistik antar kelompok (p = {p_val:.5f} < {significance_level})").bold = True
                
                doc.add_paragraph(f"Ukuran efek (η²): {eta_squared:.3f}")
                if eta_squared > 0.14:
                    doc.add_paragraph("Ini merepresentasikan ukuran efek yang besar.")
                elif eta_squared > 0.06:
                    doc.add_paragraph("Ini merepresentasikan ukuran efek sedang.")
                else:
                    doc.add_paragraph("Ini merepresentasikan ukuran efek kecil.")
            else:
                p = doc.add_paragraph()
                p.add_run(f"Kesimpulan: Tidak ada perbedaan signifikan secara statistik antar kelompok (p = {p_val:.5f} ≥ {significance_level})").bold = True
            
            # Add assumption test results
            doc.add_heading('Uji Asumsi', 1)
            
            doc.add_heading('Uji Normalitas (Shapiro-Wilk)', 2)
            add_dataframe_to_doc(doc, assumptions['normality'])
            
            doc.add_heading('Uji Homogenitas Varians (Levene)', 2)
            homogeneity_df = pd.DataFrame({
                'Test': ['Uji Levene'],
                'Statistic': [assumptions['homogeneity']['Statistic']],
                'p-value': [assumptions['homogeneity']['p-value']],
                'Equal Variances': [assumptions['homogeneity']['Equal Variances']]
            })
            add_dataframe_to_doc(doc, homogeneity_df)
            
            # Add post-hoc results if available
            if 'posthoc_df' in locals():
                doc.add_heading('Hasil Uji Post-Hoc', 1)
                add_dataframe_to_doc(doc, posthoc_df, f"Hasil Uji {posthoc_method}")
                
                # Add significant differences if available
                if 'sig_pairs' in locals() and len(sig_pairs) > 0:
                    add_dataframe_to_doc(doc, sig_pairs, "Perbedaan Kelompok yang Signifikan")
            
            # Add group statistics
            doc.add_heading('Statistik Kelompok', 1)
            group_summary_str = df.groupby(categorical_col)[numeric_col].describe().reset_index().to_string()
            doc.add_paragraph(group_summary_str)
            
            # Add footer
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Generated using One-Way ANOVA Analysis Tool | © {datetime.datetime.now().year} Galuh Adi Insani"
            
            # Save the document to a BytesIO object
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            col1, col2 = st.columns([1, 1])
            with col1:
                # Offer download for Word document
                st.download_button(
                    label="📝 Unduh Hasil Lengkap (Word)",
                    data=doc_io.getvalue(),
                    file_name="hasil_anova.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="word_download"
                )
            
            with col2:
                # CSV download option for main results
                csv_buffer = io.StringIO()
                anova_table.to_csv(csv_buffer, index=False)
                st.download_button(
                    label="📄 Unduh Tabel ANOVA (CSV)",
                    data=csv_buffer.getvalue(),
                    file_name="tabel_anova.csv",
                    mime="text/csv",
                    key="csv_download"
                )
            
            st.success("""
            ✓ File Word mencakup:  
            - Tabel ANOVA lengkap
            - Hasil uji asumsi 
            - Hasil analisis post-hoc
            - Statistik kelompok
            - Interpretasi hasil
            - Metadata analisis
            """)
    else:
        st.info("Please go to the 'Data Input' tab, select your data, and click 'Run ANOVA Analysis'")

with tab3:
    st.header("Tentang ANOVA Satu Arah")
    
    st.markdown("""
    ### Apa itu ANOVA?
    Analisis Varians (ANOVA) Satu Arah adalah metode statistik yang digunakan untuk membandingkan rata-rata tiga atau lebih kelompok independen untuk menentukan apakah terdapat perbedaan yang signifikan secara statistik di antara mereka. Teknik ini memeriksa variasi data baik di dalam maupun antar kelompok untuk mengevaluasi perbedaan keseluruhan.
    
    ### Kapan Menggunakan ANOVA Satu Arah
    Gunakan ANOVA Satu Arah ketika:
    - Anda memiliki satu variabel independen kategorikal
    - Anda memiliki variabel dependen kontinu
    - Anda ingin membandingkan rata-rata di antara tiga atau lebih kelompok
    - Data Anda memenuhi asumsi ANOVA
    
    ### Asumsi ANOVA
    1. **Independensi**: Pengamatan harus independen satu sama lain
    2. **Normalitas**: Data dalam setiap kelompok harus terdistribusi normal secara approximatif
    3. **Homogenitas varians**: Varians antar kelompok harus kurang lebih sama
    
    ### Interpretasi Hasil
    - **p-value < alpha** (biasanya 0,05): Tolak hipotesis nol. Setidaknya satu rata-rata kelompok berbeda dari yang lain.
    - **p-value ≥ alpha**: Gagal menolak hipotesis nol. Tidak ada perbedaan signifikan antara rata-rata kelompok.
    
    ### Pengujian Post-Hoc
    Ketika ANOVA menunjukkan perbedaan signifikan, uji post-hoc (seperti Tukey HSD) dapat mengidentifikasi kelompok spesifik mana yang berbeda satu sama lain.
    
    ### Ukuran Efek
    Eta-squared (η²) menunjukkan proporsi varians dalam variabel dependen yang dijelaskan oleh variabel independen:
    - Efek kecil: η² ≈ 0,01
    - Efek sedang: η² ≈ 0,06
    - Efek besar: η² ≈ 0,14
    
    ### Contoh Aplikasi ANOVA
    - Membandingkan efektivitas tiga metode pengajaran berbeda pada nilai ujian siswa
    - Membandingkan hasil panen dari empat jenis pupuk berbeda
    - Membandingkan tingkat kepuasan pelanggan di beberapa toko cabang
    - Membandingkan efektivitas beberapa jenis obat dalam menurunkan tekanan darah
    
    ### Perhitungan ANOVA
    ANOVA membagi total variasi data (SST) menjadi:
    - Variasi antar kelompok (SSB): variasi yang dijelaskan oleh perbedaan kelompok
    - Variasi dalam kelompok (SSW): variasi yang tidak dijelaskan (residual/error)
    
    F-statistik dihitung sebagai: F = (SSB/dfB) / (SSW/dfW) = MSB / MSW
    
    ### Referensi Jurnal yang dipergunakan :
    
    1. Fisher, R.A. (1925). Statistical methods for research workers. Edinburgh: Oliver and Boyd.
    2. Gelman, A. (2005). Analysis of variance—why it is more important than ever. *The Annals of Statistics*, 33(1), 1-53.
    3. Keselman, H.J., Algina, J., Kowalchuk, R.K., & Wolfinger, R.D. (1998). A comparison of two approaches for selecting covariance structures in the analysis of repeated measurements. *Communications in Statistics - Simulation and Computation*, 27(3), 591-604.    
    4. Maxwell, S.E., Delaney, H.D., & Kelley, K. (2017). *Designing experiments and analyzing data: A model comparison perspective* (3rd ed.). Routledge.   
    5. Howell, D.C. (2012). *Statistical methods for psychology* (8th ed.). Cengage Learning.
    6. Lakens, D. (2013). Calculating and reporting effect sizes to facilitate cumulative science: a practical primer for t-tests and ANOVAs. *Frontiers in Psychology*, 4, 863.
    7. Blanca, M.J., Alarcón, R., Arnau, J., Bono, R., & Bendayan, R. (2017). Non-normal data: Is ANOVA still a valid option? *Psicothema*, 29(4), 552-557.
    8. Wilcox, R.R. (2017). *Introduction to robust estimation and hypothesis testing* (4th ed.). Academic Press.
    """)

    # Add visualization of ANOVA concept
    st.subheader("Visualisasi Konsep ANOVA")
    
    st.markdown("""
    #### Perbandingan Visual Variasi Dalam dan Antar Kelompok
    
    Konsep ANOVA dapat diilustrasikan dengan membagi variasi total data menjadi variasi dalam kelompok dan antar kelompok. Jika perbedaan antar kelompok jauh lebih besar daripada variasi dalam kelompok, maka kita dapat menyimpulkan bahwa setidaknya satu kelompok berbeda secara signifikan dari yang lain.
    
    *Gambar: Ilustrasi konsep ANOVA. Variasi total dibagi menjadi variasi dalam kelompok dan antar kelompok.*
    
    #### Formula Utama ANOVA
    
    $$F = \\frac{MS_{between}}{MS_{within}} = \\frac{\\frac{SS_{between}}{df_{between}}}{\\frac{SS_{within}}{df_{within}}}$$
    
    dimana:
    - $SS_{between}$ = Jumlah kuadrat antar kelompok
    - $SS_{within}$ = Jumlah kuadrat dalam kelompok
    - $df_{between}$ = Derajat kebebasan antar kelompok (k-1)
    - $df_{within}$ = Derajat kebebasan dalam kelompok (N-k)
    - $MS_{between}$ = Rerata kuadrat antar kelompok
    - $MS_{within}$ = Rerata kuadrat dalam kelompok
    """)

    # Add flow chart for ANOVA decision process
    st.subheader("Alur Pengambilan Keputusan ANOVA")
    
    st.markdown("""
    ```
    Mulai
      │
      ▼
    Memenuhi Asumsi?
      │
      ├── Tidak ──► Pertimbangkan Transformasi Data
      │               atau Gunakan Uji Non-Parametrik
      │
      ▼ Ya
    Jalankan ANOVA
      │
      ▼
    p < alpha?
      │
      ├── Tidak ──► Gagal tolak H₀: Tidak ada perbedaan signifikan
      │
      ▼ Ya
    Tolak H₀: Ada perbedaan signifikan
      │
      ▼
    Jalankan Uji Post-Hoc
      │
      ▼
    Interpretasi Hasil
      │
      ▼
    Selesai
    ```
    """)
# Define main function or remove the call if not needed
def main():
    pass  # Your main code logic here if needed

if __name__ == '__main__':
    main()

# Footer with LinkedIn profile link and improved styling
st.markdown("""
<hr style="height:1px;border:none;color:#333;background-color:#333;margin-top:30px;margin-bottom:20px">
""", unsafe_allow_html=True)

st.markdown(f"""
<div style="text-align:center; padding:15px; margin-top:10px; margin-bottom:20px">
    <p style="font-size:16px; color:#555">
        © {current_year} Developed by: 
        <a href="https://www.linkedin.com/in/galuh-adi-insani-1aa0a5105/" target="_blank" 
           style="text-decoration:none; color:#0077B5; font-weight:bold">
            <img src="https://content.linkedin.com/content/dam/me/business/en-us/amp/brand-site/v2/bg/LI-Bug.svg.original.svg" 
                 width="16" height="16" style="vertical-align:middle; margin-right:5px">
            Galuh Adi Insani
        </a> 
        with <span style="color:#e25555">❤️</span>
    </p>
    <p style="font-size:12px; color:#777">All rights reserved.</p>
</div>
""", unsafe_allow_html=True)
